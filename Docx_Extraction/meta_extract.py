
# -*- coding: utf-8 -*-
"""
Created on Fri Aug  2 13:42:08 2019

@author: 80054
"""

import re
import os
import docx
import json
import shutil
import zipfile
from bs4 import BeautifulSoup 
from win32com.client import Dispatch
import win32com.client as win32
from nltk.stem.porter import PorterStemmer
ps = PorterStemmer()



def get_unzip_final(path):
        final_copy_file_path=(os.path.splitext(path)[0].split('Sample')[0]+"\\"+"Final_copy.docx")     
        shutil.copy2(path,final_copy_file_path)
        final_copy_file_zip_path=final_copy_file_path.replace('docx','zip')
        os.replace(final_copy_file_path,final_copy_file_zip_path)
        final_copy_file_Unzip_path=os.path.splitext(path)[0].split('Sample')[0]+"\\"+"Final"
        zip_img = zipfile.ZipFile(final_copy_file_zip_path, 'r')
        zip_img.extractall(final_copy_file_Unzip_path)
        return final_copy_file_Unzip_path,final_copy_file_zip_path




list_kywd = ["graphical abstract","highlights","abstract","keywords","nomenclature","introduction","disclosure","author contribution","funding source","ethical approval","conflict of interest","acknowledgement","appendix","references","biographies","glossary","schemes"]
stem_dict={'graphical abstract':"graphical abstract",'highlight':"Highlights",'abstract':"Abstract",'keyword':"keywords",'nomenclatur':"Nomenclature",'introduct':"Introduction",'disclosur':"Disclosure",'author contribut':"Author contribution",'funding sourc':"Funding source",'ethical approv':"Ethical approval",'conflict of interest':"Conflict of interest",'acknowledg':"Acknowledgement",'appendix':"Appendix",'refer':"References",'biographi':"Biographies",'glossari':"Glossary",'scheme':"Schemes"}

list_kywd_stem=[ps.stem(word) for word in list_kywd] 


def readtxt(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text.lower())
    return fullText


def is_grey_scale(img_path):
	from PIL import Image
	img = Image.open(img_path).convert('RGB')
	w,h = img.size
	for i in range(w):
		for j in range(h):
			r,g,b = img.getpixel((i,j))
			if r != g != b: return "color image"
	return "b/w image"


def get_head_para_no(str1):
    head_para_no=[]
    kw_para_dict={}
    for i in range(len(str1)):
        str1_lower=str1[i].lower()
        s=re.sub(r"\b\d+\b", " ", str1_lower)
        s1 = re.sub(r'[?|$|.|:|,|!]',r'',s)
        s2=s1.split()
        if len(s2)<=3 and len(s2)!=0:
            s3=ps.stem(' '.join(s2))
            for a in range(len(list_kywd_stem)):
                if list_kywd_stem[a] in s3:
                    kw_para_dict.update({list_kywd_stem[a]:i})
                    head_para_no.append(i)
    return head_para_no,kw_para_dict


def page_count(path):
    from win32com.client import Dispatch
    word = Dispatch('Word.Application')
    word.Visible = False
    word = word.Documents.Open(path)
    word.Repaginate()
    num_of_sheets = word.ComputeStatistics(2)
    return num_of_sheets
    word.close

def get_table_count(path):
    import win32com.client as win32
    word = win32.Dispatch("Word.Application")
    word.Visible = 0
    word.Documents.Open(path)
    doc = word.ActiveDocument
    table_cnt=doc.Tables.Count
    return table_cnt   


def auth_name(str1):
    import re
    auth_para=[]
    for i in str1:
        if i!='' :
            auth_para.append(i)
    s1 = re.sub(r'[?|$||:|!|*|§]|\d+|',r'',auth_para[1])
    auth_name_str = re.sub(',,',',',s1)
    title=auth_para[0]
    return title,auth_name_str


def get_contents(path1):
    path=r'C:\Users\80054\Desktop\uniproof\sample_jid_aid'+"\\"+path1+".docx"
    
    try:
        str1=(readtxt(path))
    except:
        print("Exception in reading docx file")
     
    try:
        para_no_list,head_kywd=get_head_para_no(str1)
    except:
        print("Exception in getting sub_head para no")
    
     
    listx=[]               
    for i in range(len(para_no_list)-1):  
        listy=[]  
        for j in range(para_no_list[i],para_no_list[i+1]):  
            if str1[j]!='':
                listy.append(str1[j])
        listx.append(listy)
    listx.append(str1[j+1:])
    
    
    data = {}
    for i in range(len(listx)):
        data[listx[i][0]] = listx[i][1:]
        json_data = json.dumps(data)
    
    final_copy_file_Unzip_path,final_copy_file_zip_path=get_unzip_final(path)
    img_cnt_dict={"Total_img_cnt":0,"Total black/white img":0}
    for root, dir_list, file_list in os.walk(final_copy_file_Unzip_path):
        img_cnt=0
        col_cnt=0
        b_w_cnt=0
        for file_name in file_list:
            img_cnt=img_cnt+1
            if 'image' in file_name:
                image_path=os.path.join(root,file_name)
                is_grey_test=is_grey_scale(image_path)
                img_cnt_dict.update({"Total_img_cnt":img_cnt})
                if is_grey_test=="color image":
                    col_cnt=col_cnt+1
                    img_cnt_dict.update({"Total color img":col_cnt})
                else:
                    b_w_cnt=b_w_cnt+1
                    img_cnt_dict.update({"Total black/white img":b_w_cnt})
            if 'app.xml' in file_name:
                app_xml_path=os.path.join(root,file_name)
    with open(app_xml_path,'r') as file:
        text=file.read()
    soup = BeautifulSoup(text,'xml')
    page_cnt = soup.find('Pages').text
    word_cnt = soup.find('Words').text
    

    from docx.api import Document

    # Load the first table from your document. In your example file,
    # there is only one table, so I just grab the first one.
    document = Document(path)
    table = document.tables
    table_cnt=len(table)


    Title,Auth_name=auth_name(str1)

    
    img_cnt_dict.update({"Total_page":page_cnt,"Total_word":word_cnt,"Total table":table_cnt,"Authors_name":Auth_name,"Title":Title})
    data.update(img_cnt_dict)
    with open(r'C:\Users\80054\Desktop\uniproof\json_log\\'+path1+'.json', 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
    
    
    shutil.rmtree(final_copy_file_Unzip_path,ignore_errors=True)
    os.remove(final_copy_file_zip_path)
    return data


get_contents('Sample_1')







